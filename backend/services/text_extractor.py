"""Text Extraction Utility for Knowledge Spaces

Extracts raw text from various file formats for RAG indexing.

Supported formats:
- PDF (application/pdf)
- DOCX (application/vnd.openxmlformats-officedocument.wordprocessingml.document)
- TXT (text/plain)
- Markdown (text/markdown)
- HTML (text/html)

Usage:
    from services.text_extractor import extract_text
    
    text = extract_text("/path/to/file.pdf", "application/pdf")
    print(f"Extracted {len(text)} characters")

Author: DRS Implementation Team
Spec: §17 Knowledge Spaces, Task 2.1
"""

import logging
import re
from pathlib import Path
from typing import Optional

# PDF extraction
try:
    from PyPDF2 import PdfReader
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    logging.warning("PyPDF2 not installed. PDF extraction will fail.")

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

# DOCX extraction
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not installed. DOCX extraction will fail.")

# HTML extraction
try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False
    logging.warning("beautifulsoup4 not installed. HTML extraction will be basic.")


logger = logging.getLogger(__name__)


class TextExtractionError(Exception):
    """Raised when text extraction fails."""
    pass


def extract_text(file_path: str, mime_type: str) -> str:
    """Extract text from file based on MIME type.
    
    Args:
        file_path: Path to file on disk
        mime_type: MIME type (e.g., 'application/pdf')
    
    Returns:
        Extracted text (normalized)
    
    Raises:
        TextExtractionError: If extraction fails
        FileNotFoundError: If file doesn't exist
    """
    file_path_obj = Path(file_path)
    
    if not file_path_obj.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    logger.info(f"Extracting text from {file_path_obj.name} ({mime_type})")
    
    try:
        # Route to appropriate extractor
        if mime_type == "application/pdf":
            text = _extract_pdf(file_path_obj)
        elif mime_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ]:
            text = _extract_docx(file_path_obj)
        elif mime_type in ["text/plain", "text/markdown"]:
            text = _extract_text_file(file_path_obj)
        elif mime_type == "text/html":
            text = _extract_html(file_path_obj)
        else:
            # Fallback: try as plain text
            logger.warning(f"Unknown MIME type {mime_type}, trying as text")
            text = _extract_text_file(file_path_obj)
        
        # Normalize
        text = _normalize_text(text)
        
        logger.info(f"Extracted {len(text)} characters from {file_path_obj.name}")
        return text
    
    except Exception as e:
        logger.error(f"Failed to extract text from {file_path}: {e}")
        raise TextExtractionError(f"Extraction failed: {e}") from e


def _extract_pdf(file_path: Path) -> str:
    """Extract text from PDF using PyPDF2 or pdfplumber."""
    if not PYPDF2_AVAILABLE and not PDFPLUMBER_AVAILABLE:
        raise TextExtractionError(
            "Neither PyPDF2 nor pdfplumber installed. Install with: "
            "pip install PyPDF2 pdfplumber"
        )
    
    # Try PyPDF2 first (faster)
    if PYPDF2_AVAILABLE:
        try:
            reader = PdfReader(str(file_path))
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            
            if text.strip():  # Success
                return text
            else:
                logger.warning("PyPDF2 extracted empty text, trying pdfplumber")
        except Exception as e:
            logger.warning(f"PyPDF2 failed: {e}, trying pdfplumber")
    
    # Fallback to pdfplumber (more robust but slower)
    if PDFPLUMBER_AVAILABLE:
        with pdfplumber.open(file_path) as pdf:
            text = ""
            for page in pdf.pages:
                text += page.extract_text() or ""
            return text
    
    raise TextExtractionError("PDF extraction failed with all methods")


def _extract_docx(file_path: Path) -> str:
    """Extract text from DOCX using python-docx."""
    if not DOCX_AVAILABLE:
        raise TextExtractionError(
            "python-docx not installed. Install with: pip install python-docx"
        )
    
    doc = Document(str(file_path))
    
    # Extract paragraphs
    text = "\n".join([para.text for para in doc.paragraphs])
    
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join([cell.text for cell in row.cells])
            text += "\n" + row_text
    
    return text


def _extract_text_file(file_path: Path) -> str:
    """Extract text from plain text file with encoding detection."""
    # Try UTF-8 first
    try:
        return file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        logger.warning(f"UTF-8 decode failed for {file_path.name}, trying latin-1")
    
    # Fallback to latin-1 (never fails)
    try:
        return file_path.read_text(encoding="latin-1")
    except Exception as e:
        raise TextExtractionError(f"Failed to read text file: {e}") from e


def _extract_html(file_path: Path) -> str:
    """Extract text from HTML file."""
    html_content = file_path.read_text(encoding="utf-8", errors="ignore")
    
    if BS4_AVAILABLE:
        soup = BeautifulSoup(html_content, "html.parser")
        
        # Remove script and style tags
        for script in soup(["script", "style"]):
            script.extract()
        
        # Get text
        text = soup.get_text(separator="\n")
    else:
        # Basic HTML stripping (regex fallback)
        text = re.sub(r"<[^>]+>", "", html_content)
    
    return text


def _normalize_text(text: str) -> str:
    """Normalize extracted text.
    
    - Remove zero-width characters
    - Replace control characters with space
    - Collapse multiple whitespace
    - Strip leading/trailing whitespace
    """
    # Remove zero-width characters
    text = re.sub(r"[\u200b\u200c\u200d\ufeff]", "", text)
    
    # Replace control characters (except newline/tab) with space
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", " ", text)
    
    # Collapse multiple spaces (but preserve newlines)
    text = re.sub(r" +", " ", text)
    
    # Collapse multiple newlines to max 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    
    # Strip
    text = text.strip()
    
    return text


def get_supported_mime_types() -> list[str]:
    """Get list of supported MIME types based on installed libraries."""
    supported = ["text/plain", "text/markdown", "text/html"]
    
    if PYPDF2_AVAILABLE or PDFPLUMBER_AVAILABLE:
        supported.append("application/pdf")
    
    if DOCX_AVAILABLE:
        supported.extend([
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ])
    
    return supported


if __name__ == "__main__":
    # Test extraction with sample files
    import sys
    
    if len(sys.argv) < 3:
        print("Usage: python text_extractor.py <file_path> <mime_type>")
        print("\nSupported MIME types:")
        for mime in get_supported_mime_types():
            print(f"  - {mime}")
        sys.exit(1)
    
    file_path = sys.argv[1]
    mime_type = sys.argv[2]
    
    try:
        text = extract_text(file_path, mime_type)
        print("=" * 80)
        print(f"Extracted {len(text)} characters from {Path(file_path).name}")
        print("=" * 80)
        print(text[:500])  # First 500 chars
        if len(text) > 500:
            print("\n[...]\n")
            print(text[-200:])  # Last 200 chars
    except TextExtractionError as e:
        print(f"❌ Error: {e}")
        sys.exit(1)
