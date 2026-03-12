"""Export service for generating PDF, DOCX, and Markdown documents."""

from enum import Enum
from io import BytesIO
from typing import List
from datetime import datetime

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib import colors

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from database.models import Run, Section


class CitationStyle(str, Enum):
    """Supported citation styles."""
    APA = "APA"
    MLA = "MLA"
    CHICAGO = "Chicago"


class PDFExporter:
    """Export research documents to PDF using ReportLab."""
    
    def __init__(self, citation_style: CitationStyle = CitationStyle.APA):
        self.citation_style = citation_style
        self.styles = getSampleStyleSheet()
        self._setup_styles()
    
    def _setup_styles(self):
        """Configure custom paragraph styles."""
        # Title style
        self.styles.add(ParagraphStyle(
            name='Title',
            parent=self.styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=30,
            alignment=TA_CENTER,
        ))
        
        # Section heading
        self.styles.add(ParagraphStyle(
            name='SectionHeading',
            parent=self.styles['Heading1'],
            fontSize=16,
            textColor=colors.HexColor('#2563eb'),
            spaceBefore=20,
            spaceAfter=12,
        ))
        
        # Body text
        self.styles.add(ParagraphStyle(
            name='BodyJustified',
            parent=self.styles['Normal'],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceAfter=12,
        ))
    
    def export(self, run: Run, sections: List[Section]) -> bytes:
        """Generate PDF document."""
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=inch,
            leftMargin=inch,
            topMargin=inch,
            bottomMargin=inch,
        )
        
        story = []
        
        # Cover page
        story.extend(self._create_cover_page(run))
        story.append(PageBreak())
        
        # Table of contents
        story.extend(self._create_toc(sections))
        story.append(PageBreak())
        
        # Sections
        for section in sections:
            story.extend(self._create_section(section))
            story.append(Spacer(1, 0.2 * inch))
        
        # Bibliography
        story.append(PageBreak())
        story.extend(self._create_bibliography(run))
        
        doc.build(story)
        return buffer.getvalue()
    
    def _create_cover_page(self, run: Run) -> List:
        """Generate cover page."""
        elements = []
        
        elements.append(Spacer(1, 2 * inch))
        elements.append(Paragraph(run.topic, self.styles['Title']))
        elements.append(Spacer(1, 0.5 * inch))
        
        # Metadata table
        metadata = [
            ['Generated', datetime.now().strftime('%B %d, %Y')],
            ['Quality Preset', run.quality_preset or 'Balanced'],
            ['Target Words', str(run.target_words or 'N/A')],
        ]
        
        table = Table(metadata, colWidths=[2*inch, 3*inch])
        table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('TEXTCOLOR', (0, 0), (0, -1), colors.grey),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ]))
        
        elements.append(table)
        return elements
    
    def _create_toc(self, sections: List[Section]) -> List:
        """Generate table of contents."""
        elements = []
        elements.append(Paragraph('Table of Contents', self.styles['Heading1']))
        elements.append(Spacer(1, 0.3 * inch))
        
        for idx, section in enumerate(sections, 1):
            toc_entry = f"{idx}. {section.title}"
            elements.append(Paragraph(toc_entry, self.styles['Normal']))
        
        return elements
    
    def _create_section(self, section: Section) -> List:
        """Generate section content."""
        elements = []
        
        # Section title
        elements.append(Paragraph(section.title, self.styles['SectionHeading']))
        
        # Content paragraphs
        paragraphs = section.content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                # Convert markdown-style citations [1] to footnotes
                formatted = self._format_citations(para)
                elements.append(Paragraph(formatted, self.styles['BodyJustified']))
        
        return elements
    
    def _format_citations(self, text: str) -> str:
        """Format inline citations."""
        # TODO: Parse [1] style citations and convert to superscript
        return text
    
    def _create_bibliography(self, run: Run) -> List:
        """Generate bibliography section."""
        elements = []
        elements.append(Paragraph('References', self.styles['Heading1']))
        elements.append(Spacer(1, 0.3 * inch))
        
        # TODO: Extract sources from run metadata and format
        elements.append(Paragraph(
            'Source references will be listed here based on citation style.',
            self.styles['Normal']
        ))
        
        return elements


class DOCXExporter:
    """Export research documents to DOCX using python-docx."""
    
    def __init__(self, citation_style: CitationStyle = CitationStyle.APA):
        self.citation_style = citation_style
    
    def export(self, run: Run, sections: List[Section]) -> bytes:
        """Generate DOCX document."""
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = run.topic
        doc.core_properties.author = 'Deep Research System'
        doc.core_properties.created = datetime.now()
        
        # Cover page
        self._add_cover_page(doc, run)
        doc.add_page_break()
        
        # Sections
        for section in sections:
            self._add_section(doc, section)
        
        # Bibliography
        doc.add_page_break()
        self._add_bibliography(doc, run)
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    
    def _add_cover_page(self, doc: Document, run: Run):
        """Add cover page to document."""
        # Title
        title = doc.add_heading(run.topic, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        # Metadata
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph(f"Quality: {run.quality_preset or 'Balanced'}")
        doc.add_paragraph(f"Words: {run.target_words or 'N/A'}")
    
    def _add_section(self, doc: Document, section: Section):
        """Add section to document."""
        # Section heading
        doc.add_heading(section.title, level=1)
        
        # Content
        paragraphs = section.content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph(para)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def _add_bibliography(self, doc: Document, run: Run):
        """Add bibliography section."""
        doc.add_heading('References', level=1)
        doc.add_paragraph('Source references based on citation style.')


class MarkdownExporter:
    """Export research documents to Markdown."""
    
    def export(self, run: Run, sections: List[Section]) -> str:
        """Generate Markdown document."""
        lines = []
        
        # Title
        lines.append(f"# {run.topic}")
        lines.append("")
        
        # Metadata
        lines.append(f"**Generated:** {datetime.now().strftime('%B %d, %Y')}")
        lines.append(f"**Quality:** {run.quality_preset or 'Balanced'}")
        lines.append(f"**Words:** {run.target_words or 'N/A'}")
        lines.append("")
        lines.append("---")
        lines.append("")
        
        # Table of contents
        lines.append("## Table of Contents")
        lines.append("")
        for idx, section in enumerate(sections, 1):
            anchor = section.title.lower().replace(' ', '-')
            lines.append(f"{idx}. [{section.title}](#{anchor})")
        lines.append("")
        lines.append("---")
        lines.append("")
        
        # Sections
        for section in sections:
            lines.append(f"## {section.title}")
            lines.append("")
            lines.append(section.content)
            lines.append("")
            lines.append("---")
            lines.append("")
        
        # References
        lines.append("## References")
        lines.append("")
        lines.append("Source references will be listed here.")
        lines.append("")
        
        return "\n".join(lines)


def format_citation_apa(title: str, author: str, year: int, url: str = None) -> str:
    """Format citation in APA style."""
    citation = f"{author} ({year}). {title}."
    if url:
        citation += f" {url}"
    return citation


def format_citation_mla(title: str, author: str, year: int, url: str = None) -> str:
    """Format citation in MLA style."""
    citation = f'{author}. "{title}." {year}.'
    if url:
        citation += f" {url}."
    return citation


def format_citation_chicago(title: str, author: str, year: int, url: str = None) -> str:
    """Format citation in Chicago style."""
    citation = f"{author}. {title}. {year}."
    if url:
        citation += f" {url}."
    return citation
