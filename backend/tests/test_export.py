"""Tests for export service."""

import pytest
from io import BytesIO
from datetime import datetime

from docx import Document

from services.export_service import (
    PDFExporter,
    DOCXExporter,
    MarkdownExporter,
    CitationStyle,
    format_citation_apa,
    format_citation_mla,
    format_citation_chicago,
)
from database.models import Run, Section


@pytest.fixture
def sample_run():
    return Run(
        doc_id="test_123",
        topic="Machine Learning Fundamentals",
        quality_preset="Balanced",
        target_words=3000,
        status="complete",
        created_at=datetime.now(),
    )


@pytest.fixture
def sample_sections():
    return [
        Section(
            doc_id="test_123",
            section_idx=0,
            title="Introduction",
            content="Machine learning has revolutionized technology. [1]\n\nThe field continues to grow. [2]",
            status="complete",
        ),
        Section(
            doc_id="test_123",
            section_idx=1,
            title="Background",
            content="Neural networks are powerful models. [3]\n\nDeep learning enables complex tasks. [4]",
            status="complete",
        ),
    ]


class TestPDFExporter:
    def test_export_generates_pdf(self, sample_run, sample_sections):
        exporter = PDFExporter(citation_style=CitationStyle.APA)
        pdf_bytes = exporter.export(sample_run, sample_sections)
        
        # Verify PDF magic bytes
        assert pdf_bytes.startswith(b'%PDF')
        assert len(pdf_bytes) > 1000
    
    def test_export_with_different_citation_styles(self, sample_run, sample_sections):
        for style in [CitationStyle.APA, CitationStyle.MLA, CitationStyle.CHICAGO]:
            exporter = PDFExporter(citation_style=style)
            pdf_bytes = exporter.export(sample_run, sample_sections)
            assert pdf_bytes.startswith(b'%PDF')


class TestDOCXExporter:
    def test_export_generates_docx(self, sample_run, sample_sections):
        exporter = DOCXExporter(citation_style=CitationStyle.APA)
        docx_bytes = exporter.export(sample_run, sample_sections)
        
        # Verify DOCX magic bytes (ZIP archive)
        assert docx_bytes.startswith(b'PK')
        assert len(docx_bytes) > 1000
    
    def test_docx_contains_title(self, sample_run, sample_sections):
        exporter = DOCXExporter()
        docx_bytes = exporter.export(sample_run, sample_sections)
        
        # Parse DOCX
        doc = Document(BytesIO(docx_bytes))
        
        # Verify title in first paragraph
        assert doc.paragraphs[0].text == sample_run.topic
    
    def test_docx_contains_sections(self, sample_run, sample_sections):
        exporter = DOCXExporter()
        docx_bytes = exporter.export(sample_run, sample_sections)
        
        doc = Document(BytesIO(docx_bytes))
        full_text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Verify all section titles present
        for section in sample_sections:
            assert section.title in full_text


class TestMarkdownExporter:
    def test_export_generates_markdown(self, sample_run, sample_sections):
        exporter = MarkdownExporter()
        md_text = exporter.export(sample_run, sample_sections)
        
        assert isinstance(md_text, str)
        assert len(md_text) > 100
    
    def test_markdown_contains_title(self, sample_run, sample_sections):
        exporter = MarkdownExporter()
        md_text = exporter.export(sample_run, sample_sections)
        
        assert f'# {sample_run.topic}' in md_text
    
    def test_markdown_contains_sections(self, sample_run, sample_sections):
        exporter = MarkdownExporter()
        md_text = exporter.export(sample_run, sample_sections)
        
        for section in sample_sections:
            assert f'## {section.title}' in md_text
            assert section.content in md_text
    
    def test_markdown_has_toc(self, sample_run, sample_sections):
        exporter = MarkdownExporter()
        md_text = exporter.export(sample_run, sample_sections)
        
        assert '## Table of Contents' in md_text
        for section in sample_sections:
            anchor = section.title.lower().replace(' ', '-')
            assert f'[{section.title}](#{anchor})' in md_text


class TestCitationFormatting:
    def test_apa_citation(self):
        citation = format_citation_apa(
            title="Neural Networks",
            author="Smith, J.",
            year=2024,
            url="https://example.com",
        )
        
        assert "Smith, J. (2024)" in citation
        assert "Neural Networks" in citation
        assert "https://example.com" in citation
    
    def test_mla_citation(self):
        citation = format_citation_mla(
            title="Deep Learning",
            author="Doe, A.",
            year=2025,
            url="https://example.com",
        )
        
        assert 'Doe, A.' in citation
        assert '"Deep Learning."' in citation
        assert '2025' in citation
    
    def test_chicago_citation(self):
        citation = format_citation_chicago(
            title="AI Research",
            author="Johnson, B.",
            year=2023,
        )
        
        assert "Johnson, B." in citation
        assert "AI Research" in citation
        assert "2023" in citation
