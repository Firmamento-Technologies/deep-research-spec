"""Publisher node (§5.22) — format and emit final document.

Assembles all approved sections into the final document,
applies formatting, resolves citation hashes to human-readable
references, appends bibliography, and writes output files (Markdown + DOCX).
"""
from __future__ import annotations

import json
import logging
import re
from datetime import datetime, timezone
from pathlib import Path

logger = logging.getLogger(__name__)

# Optional DOCX support
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False
    logger.debug("python-docx not installed — DOCX output disabled")


def publisher_node(state: dict) -> dict:
    """Assemble and publish the final document.

    Returns:
        Partial state with ``published``, ``output_path``,
        ``output_formats``, and ``publish_metadata``.
    """
    approved = state.get("approved_sections", [])
    config = state.get("config", {})
    doc_id = state.get("doc_id", "default")
    topic = state.get("topic", config.get("topic", "Untitled Document"))
    output_formats = config.get("output_formats", ["markdown", "docx"])
    citation_style = config.get("citation_style", "APA")

    if not approved:
        logger.warning("Publisher: no approved sections to publish")
        return {"published": False}

    # ── Build unified source lookup from all approved sections ────────
    all_sources: dict[str, dict] = {}
    for section in approved:
        for src in section.get("sources", []):
            sid = src.get("source_id", "")
            if sid and sid not in all_sources:
                all_sources[sid] = src
    # Also include current_sources from state (last section, fallback)
    for src in state.get("current_sources", []):
        sid = src.get("source_id", "")
        if sid and sid not in all_sources:
            all_sources[sid] = src
    # And any citation_map already in state (from citation_manager_node)
    existing_citation_map: dict[str, str] = state.get("citation_map", {})

    # ── Build citation_map and bibliography ──────────────────────────
    citation_map: dict[str, str] = {}
    bibliography: list[str] = []
    used_source_ids: set[str] = set()

    for sid, src in all_sources.items():
        inline = _format_inline_citation(src, citation_style)
        bib_entry = _format_bibliography_entry(src, citation_style)
        citation_map[sid] = inline
        bibliography.append(bib_entry)

    # Merge any existing citation_map entries not covered
    for sid, inline in existing_citation_map.items():
        if sid not in citation_map:
            citation_map[sid] = inline

    # ── Output directory ─────────────────────────────────────────────
    output_dir = Path("output") / doc_id
    output_dir.mkdir(parents=True, exist_ok=True)
    generated_files: list[str] = []

    # ── 1. Markdown ──────────────────────────────────────────────────
    sections_md = []
    for i, section in enumerate(approved):
        title = section.get("title", f"Section {i + 1}")
        content = section.get("content", section.get("draft", ""))
        # Resolve [hash_id] placeholders to inline citations
        content = _resolve_citations(content, citation_map, used_source_ids)
        sections_md.append(f"## {title}\n\n{content}")

    document = f"# {topic}\n\n" + "\n\n---\n\n".join(sections_md)

    # Append bibliography if we have sources that were actually cited
    bib_section = _build_bibliography_section(all_sources, used_source_ids, citation_style)
    if bib_section:
        document += "\n\n---\n\n" + bib_section

    if "markdown" in output_formats:
        doc_path = output_dir / "document.md"
        doc_path.write_text(document, encoding="utf-8")
        generated_files.append(str(doc_path))

    # ── 2. DOCX ──────────────────────────────────────────────────────
    docx_path = None
    if "docx" in output_formats and _HAS_DOCX:
        docx_path = _generate_docx(output_dir, topic, approved, state,
                                    citation_map=citation_map, all_sources=all_sources,
                                    citation_style=citation_style)
        if docx_path:
            generated_files.append(str(docx_path))

    # ── 3. Metadata ──────────────────────────────────────────────────
    metadata = {
        "doc_id": doc_id,
        "topic": topic,
        "sections": len(approved),
        "total_words": len(document.split()),
        "published_at": datetime.now(timezone.utc).isoformat(),
        "quality_preset": state.get("budget", {}).get("quality_preset", "balanced"),
        "total_cost": state.get("budget", {}).get("spent_dollars", 0),
        "qa_passed": state.get("qa_passed", False),
        "output_formats": [f for f in output_formats if f != "docx" or _HAS_DOCX],
        "generated_files": generated_files,
    }

    meta_path = output_dir / "metadata.json"
    meta_path.write_text(json.dumps(metadata, indent=2), encoding="utf-8")

    logger.info(
        "Publisher: wrote %d sections (%d words, $%.2f) → %s",
        len(approved), metadata["total_words"],
        metadata["total_cost"],
        ", ".join(generated_files),
    )

    return {
        "published": True,
        "output_path": generated_files[0] if generated_files else "",
        "output_formats": generated_files,
        "publish_metadata": metadata,
    }


# ── Citation Resolution Helpers ──────────────────────────────────────────────

_CITATION_RE = re.compile(r"\[([a-f0-9]{8,})\]")


def _format_inline_citation(src: dict, style: str = "APA") -> str:
    """Format a source dict into an inline citation string."""
    authors = [a for a in src.get("authors", []) if a and a.strip()]
    year = src.get("year")
    if authors:
        parts = authors[0].split(",")[0].split()
        first_author = parts[-1] if parts else "Unknown"
    else:
        first_author = "Unknown"

    if style in ("APA", "Harvard"):
        if len(authors) > 2:
            return f"({first_author} et al., {year or 'n.d.'})"
        elif len(authors) == 2:
            second_parts = authors[1].split(",")[0].split()
            second = second_parts[-1] if second_parts else "Unknown"
            return f"({first_author} & {second}, {year or 'n.d.'})"
        return f"({first_author}, {year or 'n.d.'})"
    elif style == "Chicago":
        return f"{first_author} ({year or 'n.d.'})"
    return f"({first_author}, {year or 'n.d.'})"


def _format_bibliography_entry(src: dict, style: str = "APA") -> str:
    """Format a full bibliography entry for one source."""
    authors = src.get("authors", [])
    title = src.get("title", "Untitled")
    year = src.get("year")
    publisher = src.get("publisher", "")
    doi = src.get("doi")
    url = src.get("url")

    author_str = ", ".join(authors) if authors else "Unknown"

    entry = f"{author_str} ({year or 'n.d.'}). *{title}*."
    if publisher:
        entry += f" {publisher}."
    if doi:
        entry += f" https://doi.org/{doi}"
    elif url:
        entry += f" {url}"
    return entry


def _resolve_citations(
    text: str,
    citation_map: dict[str, str],
    used_ids: set[str],
) -> str:
    """Replace [hash_id] placeholders with human-readable inline citations."""
    def replacer(match: re.Match) -> str:
        sid = match.group(1)
        if sid in citation_map:
            used_ids.add(sid)
            return citation_map[sid]
        # Unknown hash — leave as-is but shortened
        return f"[{sid[:8]}]"

    return _CITATION_RE.sub(replacer, text)


def _build_bibliography_section(
    all_sources: dict[str, dict],
    used_ids: set[str],
    style: str = "APA",
) -> str:
    """Build a References section from cited sources."""
    if not used_ids and not all_sources:
        return ""

    # Include all sources (some may be cited, some referenced in research)
    entries = []
    for sid, src in all_sources.items():
        entries.append(_format_bibliography_entry(src, style))

    if not entries:
        return ""

    entries.sort()
    return "## References\n\n" + "\n\n".join(entries)


# ── DOCX Generator ──────────────────────────────────────────────────────────

def _generate_docx(
    output_dir: Path,
    topic: str,
    sections: list[dict],
    state: dict,
    citation_map: dict[str, str] | None = None,
    all_sources: dict[str, dict] | None = None,
    citation_style: str = "APA",
) -> Path | None:
    """Generate a styled DOCX document."""
    try:
        doc = DocxDocument()

        # ── Title page ───────────────────────────────────────────────
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(topic)
        title_run.bold = True
        title_run.font.size = Pt(28)
        title_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subtitle.add_run(
            f"Generated by Deep Research System\n"
            f"{datetime.now(timezone.utc).strftime('%B %d, %Y')}"
        )
        sub_run.font.size = Pt(12)
        sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_page_break()

        # ── Sections ─────────────────────────────────────────────────
        docx_used_ids: set[str] = set()
        for i, section in enumerate(sections):
            title = section.get("title", f"Section {i + 1}")
            content = section.get("content", section.get("draft", ""))
            if citation_map:
                content = _resolve_citations(content, citation_map, docx_used_ids)

            # Section heading
            heading = doc.add_heading(title, level=2)
            heading.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

            # Body paragraphs
            for para_text in content.split("\n\n"):
                para_text = para_text.strip()
                if not para_text:
                    continue
                p = doc.add_paragraph(para_text)
                p.paragraph_format.space_after = Pt(8)
                for run in p.runs:
                    run.font.size = Pt(11)
                    run.font.name = "Calibri"

            # Separator between sections (except last)
            if i < len(sections) - 1:
                doc.add_paragraph("─" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ── Bibliography ─────────────────────────────────────────────
        if all_sources:
            doc.add_page_break()
            ref_heading = doc.add_heading("References", level=2)
            ref_heading.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            for sid, src in sorted(all_sources.items(), key=lambda x: x[1].get("authors", [""])[0] if x[1].get("authors") else ""):
                bib_entry = _format_bibliography_entry(src, citation_style)
                p = doc.add_paragraph(bib_entry)
                p.paragraph_format.space_after = Pt(6)
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"

        # ── Save ─────────────────────────────────────────────────────
        docx_path = output_dir / "document.docx"
        doc.save(str(docx_path))
        logger.info("Publisher: DOCX written to %s", docx_path)
        return docx_path

    except Exception as exc:
        logger.warning("Publisher: DOCX generation failed: %s", exc)
        return None
